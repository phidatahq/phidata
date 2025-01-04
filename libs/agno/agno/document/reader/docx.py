import io
from pathlib import Path
from typing import List, Union

from docx import Document as DocxDocument  # type: ignore

from agno.document.base import Document
from agno.document.reader.base import Reader
from agno.utils.log import logger


class DocxReader(Reader):
    """Reader for Doc/Docx files"""

    def read(self, file: Union[Path, io.BytesIO]) -> List[Document]:
        if not file:
            raise ValueError("No file provided")

        try:
            if isinstance(file, Path):
                logger.info(f"Reading: {file}")
                docx_document = DocxDocument(file)
                doc_name = file.stem
            else:  # Handle file-like object from upload
                logger.info(f"Reading uploaded file: {file.name}")
                docx_document = DocxDocument(file)
                doc_name = file.name.split(".")[0]

            doc_content = "\n\n".join([para.text for para in docx_document.paragraphs])

            documents = [
                Document(
                    name=doc_name,
                    id=doc_name,
                    content=doc_content,
                )
            ]
            if self.chunk:
                chunked_documents = []
                for document in documents:
                    chunked_documents.extend(self.chunk_document(document))
                return chunked_documents
            return documents
        except Exception as e:
            logger.error(f"Error reading file: {e}")
            return []