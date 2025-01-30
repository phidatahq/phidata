import io
from pathlib import Path
from typing import List, Union

from agno.document.base import Document
from agno.document.reader.base import Reader
from agno.utils.log import logger

try:
    from docx import Document as DocxDocument  # type: ignore
except ImportError:
    raise ImportError("The `python-docx` package is not installed. Please install it via `pip install python-docx`.")


class DocxReader(Reader):
    """Reader for Doc/Docx files"""

    def read(self, file: Union[Path, io.BytesIO]) -> List[Document]:
        try:
            if isinstance(file, Path):
                logger.info(f"Reading: {file}")
                docx_document = DocxDocument(str(file))
                doc_name = file.stem
            else:  # Handle file-like object from upload
                logger.info(f"Reading uploaded file: {file.name}")
                docx_document = DocxDocument(file)
                doc_name = file.name.split(".")[0]

            doc_content = "\n\n".join([para.text for para in docx_document.paragraphs])

            documents = [
                Document(
                    name=doc_name,
                    id=doc_name,
                    content=doc_content,
                )
            ]
            if self.chunk:
                chunked_documents = []
                for document in documents:
                    chunked_documents.extend(self.chunk_document(document))
                return chunked_documents
            return documents
        except Exception as e:
            logger.error(f"Error reading file: {e}")
            return []
